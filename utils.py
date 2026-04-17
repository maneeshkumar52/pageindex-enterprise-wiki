"""
utils.py — Shared helpers for the PageIndex Enterprise Wiki.

Covers: configuration loading, logging setup, document parsing,
text chunking, file I/O, and export utilities.
"""

from __future__ import annotations

import hashlib
import logging
import os
import re
import textwrap
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import yaml


# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

_CONFIG_CACHE: dict[str, Any] | None = None


def load_config(path: str = "config.yaml") -> dict[str, Any]:
    """Load and cache the YAML configuration file."""
    global _CONFIG_CACHE
    if _CONFIG_CACHE is None:
        config_path = Path(path)
        if not config_path.exists():
            raise FileNotFoundError(f"Configuration file not found: {path}")
        with open(config_path, "r", encoding="utf-8") as fh:
            _CONFIG_CACHE = yaml.safe_load(fh)
    return _CONFIG_CACHE


def get_nested(cfg: dict, *keys: str, default: Any = None) -> Any:
    """Safely traverse nested dict keys."""
    for k in keys:
        if isinstance(cfg, dict):
            cfg = cfg.get(k, default)
        else:
            return default
    return cfg


# ---------------------------------------------------------------------------
# Logging
# ---------------------------------------------------------------------------

def setup_logging(cfg: dict[str, Any] | None = None) -> logging.Logger:
    """Configure the application-wide logger from config."""
    if cfg is None:
        cfg = load_config()
    log_cfg = cfg.get("logging", {})
    log_level = getattr(logging, log_cfg.get("level", "INFO").upper(), logging.INFO)
    log_format = log_cfg.get("format", "%(asctime)s | %(levelname)-8s | %(name)s | %(message)s")
    log_file = log_cfg.get("file")

    logger = logging.getLogger("pageindex_wiki")
    if logger.handlers:
        return logger

    logger.setLevel(log_level)
    formatter = logging.Formatter(log_format)

    console = logging.StreamHandler()
    console.setLevel(log_level)
    console.setFormatter(formatter)
    logger.addHandler(console)

    if log_file:
        Path(log_file).parent.mkdir(parents=True, exist_ok=True)
        file_handler = logging.FileHandler(log_file, encoding="utf-8")
        file_handler.setLevel(log_level)
        file_handler.setFormatter(formatter)
        logger.addHandler(file_handler)

    return logger


# ---------------------------------------------------------------------------
# Directory helpers
# ---------------------------------------------------------------------------

def ensure_dirs(cfg: dict[str, Any] | None = None) -> None:
    """Create all required directories from config."""
    if cfg is None:
        cfg = load_config()
    dirs = [
        get_nested(cfg, "storage", "upload_dir", default="data/uploads"),
        get_nested(cfg, "storage", "index_dir", default="data/indexes"),
        get_nested(cfg, "storage", "output_dir", default="outputs"),
        get_nested(cfg, "chromadb", "persist_directory", default="data/chromadb"),
        get_nested(cfg, "pageindex", "workspace", default="data/workspace"),
    ]
    for d in dirs:
        Path(d).mkdir(parents=True, exist_ok=True)


# ---------------------------------------------------------------------------
# Document parsing
# ---------------------------------------------------------------------------

def extract_text_from_pdf(file_path: str) -> str:
    """Extract raw text from a PDF using PyMuPDF."""
    import fitz  # PyMuPDF

    text_parts: list[str] = []
    with fitz.open(file_path) as doc:
        for page in doc:
            text_parts.append(page.get_text())
    return "\n".join(text_parts)


def extract_text_from_docx(file_path: str) -> str:
    """Extract raw text from a DOCX file."""
    from docx import Document as DocxDocument

    doc = DocxDocument(file_path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def extract_text_from_txt(file_path: str) -> str:
    """Read plain-text or Markdown files."""
    with open(file_path, "r", encoding="utf-8", errors="replace") as fh:
        return fh.read()


def extract_text(file_path: str) -> str:
    """Dispatch to the appropriate extractor based on file extension."""
    ext = Path(file_path).suffix.lower()
    extractors = {
        ".pdf": extract_text_from_pdf,
        ".docx": extract_text_from_docx,
        ".txt": extract_text_from_txt,
        ".md": extract_text_from_txt,
        ".markdown": extract_text_from_txt,
    }
    extractor = extractors.get(ext)
    if extractor is None:
        raise ValueError(f"Unsupported file type: {ext}")
    return extractor(file_path)


# ---------------------------------------------------------------------------
# Text chunking
# ---------------------------------------------------------------------------

def chunk_text(
    text: str,
    chunk_size: int = 512,
    chunk_overlap: int = 64,
) -> list[dict[str, Any]]:
    """
    Split text into overlapping chunks of approximately *chunk_size* tokens.

    Returns a list of dicts: {"index": int, "text": str}
    """
    words = text.split()
    chunks: list[dict[str, Any]] = []
    start = 0
    idx = 0
    while start < len(words):
        end = min(start + chunk_size, len(words))
        chunk_words = words[start:end]
        chunks.append({"index": idx, "text": " ".join(chunk_words)})
        idx += 1
        start += chunk_size - chunk_overlap
    return chunks


# ---------------------------------------------------------------------------
# File identity
# ---------------------------------------------------------------------------

def file_hash(file_path: str) -> str:
    """Return a SHA-256 hex-digest for a file (used to detect duplicates)."""
    h = hashlib.sha256()
    with open(file_path, "rb") as fh:
        for block in iter(lambda: fh.read(1 << 16), b""):
            h.update(block)
    return h.hexdigest()


# ---------------------------------------------------------------------------
# Metadata helpers
# ---------------------------------------------------------------------------

def build_doc_metadata(file_path: str) -> dict[str, Any]:
    """Create a metadata dict for a newly uploaded document."""
    p = Path(file_path)
    return {
        "filename": p.name,
        "extension": p.suffix.lower(),
        "size_bytes": p.stat().st_size,
        "sha256": file_hash(file_path),
        "indexed_at": datetime.now(timezone.utc).isoformat(),
    }


# ---------------------------------------------------------------------------
# Export helpers
# ---------------------------------------------------------------------------

def export_to_markdown(
    query: str,
    answer: str,
    sources: list[dict[str, Any]],
    output_dir: str = "outputs",
) -> str:
    """
    Write a Q&A result to a timestamped Markdown file in *output_dir*.

    Returns the path of the created file.
    """
    Path(output_dir).mkdir(parents=True, exist_ok=True)
    ts = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    slug = re.sub(r"\W+", "_", query[:40]).strip("_").lower()
    filename = f"{ts}_{slug}.md"
    filepath = Path(output_dir) / filename

    lines = [
        f"# Query Result — {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')}",
        "",
        "## Question",
        "",
        query,
        "",
        "## Answer",
        "",
        answer,
        "",
    ]

    if sources:
        lines.append("## Sources")
        lines.append("")
        for i, src in enumerate(sources, 1):
            doc_name = src.get("filename", src.get("doc_name", "Unknown"))
            page = src.get("page", "N/A")
            lines.append(f"**{i}.** {doc_name} — page {page}")
            snippet = src.get("snippet", "")
            if snippet:
                lines.append(f"> {textwrap.shorten(snippet, width=300, placeholder='…')}")
            lines.append("")

    filepath.write_text("\n".join(lines), encoding="utf-8")
    return str(filepath)


# ---------------------------------------------------------------------------
# Markdown conversion (for DOCX → Markdown before PageIndex)
# ---------------------------------------------------------------------------

def docx_to_markdown(file_path: str, output_dir: str = "data/uploads") -> str:
    """Convert a DOCX file to Markdown and return the new path."""
    from docx import Document as DocxDocument

    doc = DocxDocument(file_path)
    md_lines: list[str] = []
    for para in doc.paragraphs:
        style = para.style.name if para.style else ""
        text = para.text.strip()
        if not text:
            continue
        if "Heading 1" in style:
            md_lines.append(f"# {text}")
        elif "Heading 2" in style:
            md_lines.append(f"## {text}")
        elif "Heading 3" in style:
            md_lines.append(f"### {text}")
        else:
            md_lines.append(text)
        md_lines.append("")

    out_name = Path(file_path).stem + ".md"
    out_path = Path(output_dir) / out_name
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text("\n".join(md_lines), encoding="utf-8")
    return str(out_path)


def txt_to_markdown(file_path: str, output_dir: str = "data/uploads") -> str:
    """Wrap a plain-text file in minimal Markdown and return the new path."""
    text = extract_text_from_txt(file_path)
    out_name = Path(file_path).stem + ".md"
    out_path = Path(output_dir) / out_name
    out_path.parent.mkdir(parents=True, exist_ok=True)
    content = f"# {Path(file_path).stem}\n\n{text}\n"
    out_path.write_text(content, encoding="utf-8")
    return str(out_path)
